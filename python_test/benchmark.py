#!/usr/bin/env python3
"""
Simple, permission-free benchmark comparing blitz_parse vs python-docx.
Now includes text extraction output for comparison.
"""

import gc
import statistics
import time
from pathlib import Path

from docx import Document

import blitz_parse

# Configuration
DOCX_PATH = "tests/assets/sample.docx"
WARMUP_RUNS = 5
BENCHMARK_RUNS = 1000
OUTLIER_THRESHOLD = 2.0


def ensure_file_exists():
    """Verify test file exists and is readable."""
    if not Path(DOCX_PATH).exists():
        raise FileNotFoundError(f"Test file not found: {DOCX_PATH}")

    file_size = Path(DOCX_PATH).stat().st_size
    print(f"Test file: {DOCX_PATH} ({file_size:,} bytes)")
    return file_size


def save_extracted_text(text, filename, method_name):
    """Save extracted text to file for comparison."""
    output_path = Path(f"extracted_text_{filename}.txt")

    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"=== EXTRACTED BY {method_name.upper()} ===\n")
            f.write(f"Length: {len(text):,} characters\n")
            f.write(f"Lines: {text.count(chr(10)) + 1}\n")
            f.write("=" * 50 + "\n\n")
            f.write(text)

        print(f"  💾 Saved to: {output_path}")
        return output_path
    except Exception as e:
        print(f"  ⚠️  Could not save {filename}: {e}")
        return None


def compare_extracted_text(text1, text2, name1, name2):
    """Compare two extracted texts and show differences."""
    print(f"\n📊 TEXT COMPARISON:")
    print(f"  {name1}: {len(text1):,} characters, {text1.count(chr(10)) + 1} lines")
    print(f"  {name2}: {len(text2):,} characters, {text2.count(chr(10)) + 1} lines")

    # Character difference
    char_diff = abs(len(text1) - len(text2))
    char_diff_pct = (
        (char_diff / max(len(text1), len(text2))) * 100
        if max(len(text1), len(text2)) > 0
        else 0
    )

    print(f"  Character difference: {char_diff:,} ({char_diff_pct:.1f}%)")

    # Line difference
    lines1 = text1.count("\n") + 1
    lines2 = text2.count("\n") + 1
    line_diff = abs(lines1 - lines2)

    print(f"  Line difference: {line_diff}")

    # Content similarity check
    if text1.strip() == text2.strip():
        print("  ✅ Content is identical (ignoring whitespace)")
    elif char_diff_pct < 5:
        print("  ✅ Content is very similar (<5% difference)")
    elif char_diff_pct < 15:
        print("  ⚠️  Content has some differences (5-15%)")
    else:
        print("  ❌ Content is significantly different (>15%)")

        # Show first difference
        min_len = min(len(text1), len(text2))
        for i in range(min_len):
            if text1[i] != text2[i]:
                context_start = max(0, i - 20)
                context_end = min(min_len, i + 20)
                print(f"  First difference at position {i}:")
                print(f"    {name1}: '{text1[context_start:context_end]}'")
                print(f"    {name2}: '{text2[context_start:context_end]}'")
                break


def benchmark_function(func, name, runs=BENCHMARK_RUNS):
    """Simple benchmark function with error tracking."""
    print(f"\nBenchmarking {name}...")

    # Warmup
    for _ in range(WARMUP_RUNS):
        gc.collect()
        try:
            func()
        except Exception as e:
            print(f"\nWarmup error: {e}")
    print(f"  Warmed up with {WARMUP_RUNS} runs")

    # Benchmark
    times = []
    errors = []
    last_result = None
    print(f"  Running {runs} iterations...", end="", flush=True)

    for i in range(runs):
        gc.collect()

        start_time = time.perf_counter()
        try:
            result = func()
            end_time = time.perf_counter()

            if not isinstance(result, str) or len(result) < 10:
                print(f"\nWarning: {name} returned suspicious result at run {i+1}")

            times.append((end_time - start_time) * 1000)
            last_result = result  # Keep the last successful result

            if i % 20 == 19:
                print(".", end="", flush=True)

        except Exception as e:
            errors.append(f"Run {i+1}: {str(e)[:50]}...")
            if i % 20 == 19:
                print("E", end="", flush=True)  # E for error

    print(" done")

    if errors:
        print(f"  ⚠️  {len(errors)} errors occurred:")
        for error in errors[:3]:  # Show first 3 errors
            print(f"     {error}")
        if len(errors) > 3:
            print(f"     ... and {len(errors) - 3} more")

    return times, last_result


def remove_outliers(times):
    """Remove statistical outliers."""
    if len(times) < 5:
        return times

    mean_time = statistics.mean(times)
    std_dev = statistics.stdev(times)

    return [t for t in times if abs(t - mean_time) <= OUTLIER_THRESHOLD * std_dev]


def print_stats(times, name):
    """Print statistics."""
    clean_times = remove_outliers(times)

    if not clean_times:
        print(f"  {name}: No valid results!")
        return None

    stats = {
        "mean": statistics.mean(clean_times),
        "median": statistics.median(clean_times),
        "min": min(clean_times),
        "max": max(clean_times),
        "count": len(clean_times),
    }

    print(f"  {name}:")
    print(f"    Median: {stats['median']:.2f} ms")
    print(f"    Mean:   {stats['mean']:.2f} ms")
    print(f"    Range:  {stats['min']:.2f} - {stats['max']:.2f} ms")
    print(f"    Valid:  {stats['count']}/{len(times)} runs")

    return stats


# Benchmark functions
def benchmark_blitz_parse():
    return blitz_parse.extract_text_py(DOCX_PATH)


def benchmark_python_docx():
    doc = Document(DOCX_PATH)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def main():
    print("=" * 60)
    print("DOCX Performance Benchmark with Text Extraction")
    print("=" * 60)

    file_size = ensure_file_exists()
    print(f"Runs: {BENCHMARK_RUNS} (after {WARMUP_RUNS} warmup)")

    # Run benchmarks and capture extracted text
    blitz_times, blitz_text = benchmark_function(
        benchmark_blitz_parse, "blitz_parse (Rust)"
    )
    python_times, python_text = benchmark_function(benchmark_python_docx, "python-docx")

    print("\n" + "=" * 60)
    print("PERFORMANCE RESULTS")
    print("=" * 60)

    blitz_stats = print_stats(blitz_times, "blitz_parse")
    python_stats = print_stats(python_times, "python-docx")

    # Save extracted text to files
    print(f"\n💾 SAVING EXTRACTED TEXT:")
    blitz_file = None
    python_file = None

    if blitz_text:
        blitz_file = save_extracted_text(
            blitz_text, "blitz_parse", "blitz_parse (Rust)"
        )

    if python_text:
        python_file = save_extracted_text(python_text, "python_docx", "python-docx")

    # Compare extracted text
    if blitz_text and python_text:
        compare_extracted_text(blitz_text, python_text, "blitz_parse", "python-docx")

    # Performance comparison
    if blitz_stats and python_stats:
        speedup = python_stats["median"] / blitz_stats["median"]
        print(f"\n🏁 PERFORMANCE SUMMARY:")
        print(f"   File size: {file_size:,} bytes")
        print(f"   blitz_parse:  {blitz_stats['median']:.2f} ms")
        print(f"   python-docx:  {python_stats['median']:.2f} ms")
        print(f"   Speedup:      {speedup:.1f}x")

        if speedup > 1:
            print(f"   🚀 blitz_parse is {speedup:.1f}x FASTER!")
        else:
            print(f"   ⚠️ blitz_parse is {1/speedup:.1f}x slower")

    # File locations
    print(f"\n📁 EXTRACTED TEXT FILES:")
    if blitz_file:
        print(f"   Rust output: {blitz_file}")
    if python_file:
        print(f"   Python output: {python_file}")

    print(f"\n💡 TIP: Compare the files with:")
    print(f"   diff extracted_text_blitz_parse.txt extracted_text_python_docx.txt")


if __name__ == "__main__":
    main()
